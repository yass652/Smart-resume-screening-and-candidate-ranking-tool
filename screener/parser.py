"""
parser.py
---------
Extracts raw text from resumes in .pdf, .docx, or .txt format.
"""

import os
import pypdf
import docx


def parse_pdf(filepath: str) -> str:
    text_chunks = []
    with open(filepath, "rb") as f:
        reader = pypdf.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks)


def parse_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    paragraphs = [p.text for p in document.paragraphs]
    # Also grab text inside tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def parse_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_resume(filepath: str) -> str:
    """Dispatches to the correct parser based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return parse_pdf(filepath)
    elif ext == ".docx":
        return parse_docx(filepath)
    elif ext == ".txt":
        return parse_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .pdf, .docx, or .txt")


def load_resumes_from_folder(folder: str) -> dict:
    """Returns {filename: raw_text} for every supported resume in a folder."""
    resumes = {}
    for fname in sorted(os.listdir(folder)):
        fpath = os.path.join(folder, fname)
        if not os.path.isfile(fpath):
            continue
        ext = os.path.splitext(fname)[1].lower()
        if ext not in (".pdf", ".docx", ".txt"):
            continue
        try:
            resumes[fname] = parse_resume(fpath)
        except Exception as e:
            print(f"[warn] Could not parse {fname}: {e}")
    return resumes
